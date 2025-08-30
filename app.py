import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import os
from openai import OpenAI
from io import StringIO
import base64
import numpy as np
from datetime import datetime, timedelta
import random

# ✅ Create OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

st.set_page_config(page_title="AI Analytics Agent", page_icon="📊", layout="wide")
st.title("📊 AI Analytics Agent")
st.write("Generate OKRs, KPIs, Funnel & Cohort Analysis in minutes.")

# Helper function to create download link for CSV
def get_csv_download_link(df, filename, text="Download CSV"):
    csv = df.to_csv(index=False)
    b64 = base64.b64encode(csv.encode()).decode()
    href = f'<a href="data:file/csv;base64,{b64}" download="{filename}">{text}</a>'
    return href

# Function to generate sample data based on business context
def generate_sample_data(business_context):
    """Generate realistic sample data based on the business context"""
    
    # Generate sample events based on business type
    if any(word in business_context.lower() for word in ['ecommerce', 'shop', 'store', 'retail']):
        events = ['page_view', 'product_view', 'add_to_cart', 'checkout_start', 'purchase_complete']
    elif any(word in business_context.lower() for word in ['saas', 'software', 'app', 'platform']):
        events = ['signup', 'login', 'feature_use', 'upgrade_plan', 'subscription_active']
    elif any(word in business_context.lower() for word in ['content', 'blog', 'media', 'news']):
        events = ['page_view', 'article_read', 'comment_post', 'share_content', 'subscribe_newsletter']
    else:
        events = ['page_view', 'signup', 'feature_use', 'conversion', 'retention_event']
    
    # Generate sample data
    np.random.seed(42)  # For reproducible results
    n_users = 1000
    n_events = 5000
    
    # Generate user IDs
    user_ids = [f"user_{i:04d}" for i in range(1, n_users + 1)]
    
    # Generate timestamps over the last 8 weeks
    end_date = datetime.now()
    start_date = end_date - timedelta(weeks=8)
    
    timestamps = []
    event_list = []
    user_list = []
    
    for _ in range(n_events):
        # Random timestamp within the date range
        random_days = random.randint(0, (end_date - start_date).days)
        random_seconds = random.randint(0, 24*60*60)
        timestamp = start_date + timedelta(days=random_days, seconds=random_seconds)
        
        # Weight events to create realistic funnel
        if random.random() < 0.4:  # 40% page views
            event = events[0]
        elif random.random() < 0.7:  # 30% of remaining (18% total) for next event
            event = events[1]
        elif random.random() < 0.8:  # 10% of remaining (8% total) for next event
            event = events[2]
        elif random.random() < 0.9:  # 10% of remaining (8% total) for next event
            event = events[3]
        else:  # 10% of remaining (6% total) for final event
            event = events[4]
        
        # Random user
        user_id = random.choice(user_ids)
        
        timestamps.append(timestamp)
        event_list.append(event)
        user_list.append(user_id)
    
    # Create DataFrame
    sample_df = pd.DataFrame({
        'user_id': user_list,
        'event': event_list,
        'timestamp': timestamps
    })
    
    return sample_df, events

# Function to create a Word document report
def create_word_report(context, okrs_kpis, sample_data_info, funnel_df, cohort_summary, retention_matrix, ai_analysis, exec_summary, funnel_fig=None, cohort_fig=None):
    """
    Creates a comprehensive Word document report.
    """
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        
        doc = Document()
        
        # Title
        title = doc.add_heading("AI Analytics Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Timestamp
        timestamp = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Business Context
        doc.add_heading("Business Context", level=1)
        doc.add_paragraph(context)
        
        # OKRs & KPIs
        doc.add_heading("OKRs & KPIs", level=1)
        doc.add_paragraph(okrs_kpis)
        
        # Sample Data Information
        if sample_data_info:
            doc.add_heading("Sample Data Overview", level=1)
            doc.add_paragraph(sample_data_info)
        
        # Funnel Analysis
        doc.add_heading("Funnel Analysis", level=1)
        if not funnel_df.empty:
            # Add funnel table
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Event'
            hdr_cells[1].text = 'Unique Users'
            hdr_cells[2].text = 'Conversion Rate (%)'
            
            for _, row in funnel_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['event'])
                row_cells[1].text = str(row['unique_users'])
                row_cells[2].text = f"{row['conversion_rate']:.2f}"
        
        # Add funnel chart if available
        if funnel_fig:
            doc.add_heading("Funnel Analysis Chart", level=2)
            # Save figure to bytes
            img_buffer = io.BytesIO()
            funnel_fig.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
            img_buffer.seek(0)
            doc.add_picture(img_buffer, width=Inches(6))
        
        # Cohort Analysis
        doc.add_heading("Cohort Analysis", level=1)
        if not cohort_summary.empty:
            # Add cohort table
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Cohort Week'
            hdr_cells[1].text = 'Week 0 Users'
            hdr_cells[2].text = 'Week 1 Retention (%)'
            hdr_cells[3].text = 'Week 2 Retention (%)'
            hdr_cells[4].text = 'Week 3 Retention (%)'
            hdr_cells[5].text = 'Week 4 Retention (%)'
            
            for _, row in cohort_summary.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['cohort_week'])
                row_cells[1].text = str(row['week_0_users'])
                row_cells[2].text = f"{row['week_1_retention']:.1f}"
                row_cells[3].text = f"{row['week_2_retention']:.1f}"
                row_cells[4].text = f"{row['week_3_retention']:.1f}"
                row_cells[5].text = f"{row['week_4_retention']:.1f}"
        
        # Add cohort chart if available
        if cohort_fig:
            doc.add_heading("Cohort Retention Chart", level=2)
            # Save figure to bytes
            img_buffer = io.BytesIO()
            cohort_fig.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
            img_buffer.seek(0)
            doc.add_picture(img_buffer, width=Inches(6))
        
        # AI Analysis
        doc.add_heading("AI-Generated Analysis", level=1)
        doc.add_paragraph(ai_analysis)
        
        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(exec_summary)
        
        # Key Metrics Summary
        doc.add_heading("Key Metrics Summary", level=1)
        
        if not funnel_df.empty:
            doc.add_paragraph("Funnel Metrics:")
            p = doc.add_paragraph()
            p.add_run(f"• Total funnel steps: {len(funnel_df)}")
            p.add_run(f"\n• Overall conversion rate: {funnel_df.iloc[-1]['conversion_rate']:.2f}%")
            if len(funnel_df) > 1:
                p.add_run(f"\n• Biggest drop-off: {funnel_df.iloc[0]['event']} to {funnel_df.iloc[1]['event']} ({(funnel_df.iloc[0]['conversion_rate'] - funnel_df.iloc[1]['conversion_rate']):.2f}% drop)")
        
        if not cohort_summary.empty:
            doc.add_paragraph("Cohort Metrics:")
            p = doc.add_paragraph()
            avg_week1 = cohort_summary['week_1_retention'].mean()
            avg_week2 = cohort_summary['week_2_retention'].mean()
            p.add_run(f"• Average Week 1 retention: {avg_week1:.1f}%")
            p.add_run(f"\n• Average Week 2 retention: {avg_week2:.1f}%")
            p.add_run(f"\n• Number of cohorts analyzed: {len(cohort_summary)}")
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except ImportError:
        st.error("python-docx library not installed. Please install it with: pip install python-docx")
        return None
    except Exception as e:
        st.error(f"Error creating Word document: {str(e)}")
        return None

# Business context input
context = st.text_area("📝 Describe your business or product")

uploaded_file = st.file_uploader(
    "📂 Upload event data (CSV with user_id, event, timestamp)", 
    type=["csv"]
)

# Add option to generate sample data
generate_sample = st.checkbox("🎲 Generate sample data based on my business description")

if st.button("Generate Insights"):
    if not context:
        st.error("Please provide a business description.")
    else:
        with st.spinner("Generating complete analytics insights..."):
            # ✅ Generate OKRs & KPIs
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a product analytics and strategy expert."},
                    {"role": "user", "content": f"Generate 3 OKRs and measurable KPIs for this business: {context}"}
                ]
            )
            st.subheader("🔑 Suggested OKRs & KPIs")
            st.markdown(response.choices[0].message.content)

            # Generate sample data if requested
            if generate_sample:
                sample_df, events = generate_sample_data(context)
                
                st.subheader("📊 Generated Sample Data")
                st.write(f"**Generated {len(sample_df)} events for {sample_df['user_id'].nunique()} users**")
                st.write(f"**Events included:** {', '.join(events)}")
                
                # Show sample of the data
                st.write("**Sample of generated data:**")
                st.dataframe(sample_df.head(10))
                
                # Create download link for sample data
                st.markdown(get_csv_download_link(sample_df, "sample_analytics_data.csv", "📥 Download Sample Data CSV"), unsafe_allow_html=True)
                
                # Use the generated data for analysis
                df = sample_df
                st.success("✅ Sample data generated!")
            elif uploaded_file:
                df = pd.read_csv(uploaded_file, parse_dates=['timestamp'])
            else:
                st.error("Please either upload a file or check 'Generate sample data' option.")
                st.stop()

            # --- Funnel Analysis ---
            st.subheader("📉 Funnel Analysis")
            top_events = df['event'].value_counts().index[:5]  # Show top 5 events
            funnel_counts = [df[df['event']==e]['user_id'].nunique() for e in top_events]
            
            # Create funnel analysis dataframe
            funnel_df = pd.DataFrame({
                'event': top_events,
                'unique_users': funnel_counts,
                'conversion_rate': [round((count/funnel_counts[0])*100, 2) for count in funnel_counts]
            })
            
            # Display funnel data
            st.write("**Funnel Analysis Results:**")
            st.dataframe(funnel_df)
            
            # Create and display funnel chart
            fig, ax = plt.subplots(figsize=(10, 6))
            bars = ax.bar(range(len(top_events)), funnel_counts, color='skyblue', alpha=0.7)
            ax.set_ylabel("Unique Users")
            ax.set_title("Funnel Drop-off Analysis")
            ax.set_xticks(range(len(top_events)))
            ax.set_xticklabels(top_events, rotation=45, ha='right')
            
            # Add value labels on bars
            for i, bar in enumerate(bars):
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + max(funnel_counts)*0.01,
                       f'{height}\n({funnel_df.iloc[i]["conversion_rate"]}%)',
                       ha='center', va='bottom')
            
            plt.tight_layout()
            st.pyplot(fig)
            
            # Download button for funnel analysis
            st.markdown(get_csv_download_link(funnel_df, "funnel_analysis.csv", "📥 Download Funnel Analysis CSV"), unsafe_allow_html=True)

            # --- Cohort Analysis ---
            st.subheader("👥 Cohort Analysis")
            
            # Create cohort analysis
            df['signup_week'] = df['timestamp'].dt.to_period('W')
            df['event_week'] = df['timestamp'].dt.to_period('W')
            
            # Get first activity week for each user (cohort)
            user_cohorts = df.groupby('user_id')['signup_week'].min().reset_index()
            user_cohorts.columns = ['user_id', 'cohort_week']
            
            # Merge cohort info with main dataframe
            df_with_cohorts = df.merge(user_cohorts, on='user_id')
            
            # Calculate period number (weeks since cohort)
            df_with_cohorts['period_number'] = (df_with_cohorts['event_week'] - df_with_cohorts['cohort_week']).apply(lambda x: x.n)
            
            # Create cohort matrix
            cohort_data = df_with_cohorts.groupby(['cohort_week', 'period_number'])['user_id'].nunique().reset_index()
            cohort_matrix = cohort_data.pivot(index='cohort_week', columns='period_number', values='user_id').fillna(0)
            
            # Calculate retention rates
            retention_matrix = cohort_matrix.div(cohort_matrix.iloc[:, 0], axis=0) * 100
            
            # Display cohort retention matrix
            st.write("**Cohort Retention Matrix (%):**")
            st.dataframe(retention_matrix.round(2))
            
            # Create cohort analysis summary
            cohort_summary = pd.DataFrame({
                'cohort_week': retention_matrix.index.astype(str),
                'week_0_users': cohort_matrix.iloc[:, 0],
                'week_1_retention': retention_matrix.iloc[:, 1] if len(retention_matrix.columns) > 1 else [0] * len(retention_matrix),
                'week_2_retention': retention_matrix.iloc[:, 2] if len(retention_matrix.columns) > 2 else [0] * len(retention_matrix),
                'week_3_retention': retention_matrix.iloc[:, 3] if len(retention_matrix.columns) > 3 else [0] * len(retention_matrix),
                'week_4_retention': retention_matrix.iloc[:, 4] if len(retention_matrix.columns) > 4 else [0] * len(retention_matrix)
            })
            
            # Display cohort summary
            st.write("**Cohort Analysis Summary:**")
            st.dataframe(cohort_summary)
            
            # Create cohort chart
            fig, ax = plt.subplots(figsize=(12, 8))
            for i, cohort in enumerate(retention_matrix.index[:5]):  # Show first 5 cohorts
                retention_rates = retention_matrix.loc[cohort]
                periods = range(len(retention_rates))
                ax.plot(periods, retention_rates, marker='o', label=f'Cohort {cohort}', linewidth=2)
            
            ax.set_xlabel('Weeks Since Cohort')
            ax.set_ylabel('Retention Rate (%)')
            ax.set_title('Cohort Retention Analysis')
            ax.legend()
            ax.grid(True, alpha=0.3)
            plt.tight_layout()
            st.pyplot(fig)
            
            # Download buttons for cohort analysis
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(get_csv_download_link(retention_matrix.reset_index(), "cohort_retention_matrix.csv", "📥 Download Retention Matrix CSV"), unsafe_allow_html=True)
            with col2:
                st.markdown(get_csv_download_link(cohort_summary, "cohort_analysis_summary.csv", "📥 Download Cohort Summary CSV"), unsafe_allow_html=True)
            
            # Additional insights
            st.subheader("💡 Key Insights")
            
            # Calculate average retention rates
            avg_retention = retention_matrix.mean()
            st.write(f"**Average retention rates:**")
            for period, rate in avg_retention.items():
                if period > 0:  # Skip week 0
                    st.write(f"  - Week {period}: {rate:.1f}%")
            
            # Identify best performing cohort
            if len(retention_matrix) > 1:
                best_cohort = retention_matrix.iloc[:, 1].idxmax()  # Best week 1 retention
                st.write(f"**Best performing cohort:** {best_cohort} (Week 1 retention: {retention_matrix.loc[best_cohort, 1]:.1f}%)")

            # --- AI Analysis Summary (Generated automatically) ---
            st.markdown("---")
            st.subheader("🤖 AI Analysis Summary")
            
            with st.spinner("Generating AI-powered insights..."):
                # Prepare data summary for AI
                funnel_summary = f"""
                Funnel Analysis:
                - Top events: {list(top_events)}
                - Conversion rates: {list(funnel_df['conversion_rate'])}
                - Total users at each step: {list(funnel_df['unique_users'])}
                """
                
                cohort_summary_text = f"""
                Cohort Analysis:
                - Number of cohorts: {len(retention_matrix)}
                - Average week 1 retention: {avg_retention.get(1, 0):.1f}%
                - Average week 2 retention: {avg_retention.get(2, 0):.1f}%
                - Best performing cohort: {best_cohort if len(retention_matrix) > 1 else 'N/A'}
                """
                
                business_context_summary = f"Business Context: {context}"
                
                # Generate AI insights
                ai_response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": """You are a senior product analyst and growth expert. Analyze the provided funnel and cohort data to generate actionable insights and recommendations. Focus on:
                        1. Key findings and patterns
                        2. Potential issues or opportunities
                        3. Specific actionable recommendations
                        4. Priority areas for improvement
                        5. Expected impact of suggested changes
                        
                        Be specific, data-driven, and provide concrete next steps."""},
                        {"role": "user", "content": f"""
                        Please analyze this analytics data and provide actionable insights:
                        
                        {business_context_summary}
                        
                        {funnel_summary}
                        
                        {cohort_summary_text}
                        
                        Provide a comprehensive analysis with specific recommendations for improvement.
                        """}
                    ]
                )
                
                st.markdown("### 📊 AI-Generated Analysis")
                st.markdown(ai_response.choices[0].message.content)
                
                # Generate executive summary
                exec_response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "You are a business analyst. Create a concise executive summary (2-3 bullet points) highlighting the most critical insights and immediate action items."},
                        {"role": "user", "content": f"Based on this analysis: {ai_response.choices[0].message.content}\n\nCreate a brief executive summary with key takeaways."}
                    ]
                )
                
                st.markdown("### 🎯 Executive Summary")
                st.markdown(exec_response.choices[0].message.content)

            # --- Generate Comprehensive Word Report ---
            st.markdown("---")
            st.subheader("📄 Download Complete Report")
            
            # Prepare sample data info
            sample_data_info = ""
            if generate_sample:
                sample_data_info = f"Generated {len(df)} events for {df['user_id'].nunique()} users. Events included: {', '.join(events)}"
            
            # Create Word document
            with st.spinner("Generating comprehensive Word report..."):
                word_doc_content = create_word_report(
                    context=context,
                    okrs_kpis=response.choices[0].message.content,
                    sample_data_info=sample_data_info,
                    funnel_df=funnel_df,
                    cohort_summary=cohort_summary,
                    retention_matrix=retention_matrix,
                    ai_analysis=ai_response.choices[0].message.content,
                    exec_summary=exec_response.choices[0].message.content,
                    funnel_fig=fig,  # Pass the stored funnel figure
                    cohort_fig=fig   # Pass the stored cohort figure
                )
                
                if word_doc_content:
                    # Create download link for Word document
                    b64 = base64.b64encode(word_doc_content).decode()
                    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="AI_Analytics_Report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx">📥 Download Complete Word Report</a>'
                    st.markdown(href, unsafe_allow_html=True)
                    st.success("✅ Word report generated successfully!")
                else:
                    st.error("❌ Failed to generate Word report. Please check if python-docx is installed.")

            st.success("🎉 Complete analytics report generated successfully!")
